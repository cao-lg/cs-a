# -*- coding: utf-8 -*-
"""
提取 初稿6.18 的 5 个项目 docx → 结构化 JSON（项目 → 任务 → 节(单元) → 子节 → 段落/表格）。
单元(section)边界判定（任一触发即新单元）：
  - 编号 一、二、三、...（顶层）
  - Heading 2 / Heading 3 描述性标题（排除 AI标记/任务/项目/导学/小结等）
子节(subsection)边界：Heading 4 / （一）（二） / 1. 2. 3. 编号
"""
import os, re, json
from docx import Document

SRC = r"D:\workbuddy\chain_supply\初稿6.18"
FILES = {
    "p1": ("项目一：开启数据驱动之旅_ 第二版.docx", "sales-project1", "项目一：开启数据驱动之旅"),
    "p2": ("项目二：商品选品数据分析-修改稿.docx", "sales-project2", "项目二：商品选品数据分析"),
    "p3": ("项目三：营销推广数据分析 样章3.0.docx", "sales-project3", "项目三：营销推广数据分析"),
    "p5": ("项目五：客户服务数据分析0501.docx", "sales-project5", "项目五：客户服务数据分析"),
    "p6": ("项目六：物流履约数据分析3.0.docx", "sales-project6", "项目六：物流履约数据分析"),
}

RE_TASK = re.compile(r'^任务\s*(\d+)\.(\d+)\s*(.*)$')
RE_SEC  = re.compile(r'^[一二三四五六七八九十]+、')
RE_SUB  = re.compile(r'^（?[一二三四五六七八九十]+）?[\.\s]')
RE_NUM  = re.compile(r'^\d+[\.\、]\s')
RE_H    = re.compile(r'^Heading\s*([1-6])', re.I)
RE_AI   = re.compile(r'^【')
SKIP_UNIT_TITLES = ('项目小结','项目自测','评价与反思','实训背景','实训数据','实训要求',
                    '学习目标','AI导学','智能体互动','项目','任务','表','图')

def hlevel(style):
    if style:
        m = RE_H.match(style)
        if m: return int(m.group(1))
    return 0

def clean(t):
    return re.sub(r'\s+',' ', t).strip()

def is_skip_title(t):
    return any(t.startswith(x) for x in SKIP_UNIT_TITLES)

def extract(path):
    doc = Document(path)
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map = {t._element: t for t in doc.tables}
    nodes = []
    for child in body.iterchildren():
        if child.tag.endswith('}p') and child in para_map:
            nodes.append(('p', para_map[child]))
        elif child.tag.endswith('}tbl') and child in tbl_map:
            nodes.append(('t', tbl_map[child]))

    tasks = []
    cur_task = None
    cur_sec = None
    cur_sub = None

    def new_task(tid, title):
        return {"id": tid, "title": title, "sections": []}
    def new_sec(title):
        return {"title": title, "subs": [], "paras": [], "tables": []}
    def new_sub(title):
        return {"title": title, "paras": [], "tables": []}
    def table_to_rows(tbl):
        return [[clean(c.text) for c in r.cells] for r in tbl.rows]

    def ensure_sec():
        nonlocal cur_task, cur_sec, cur_sub
        if cur_sec is None:
            if cur_task is None:
                cur_task = new_task("0.0","导论"); tasks.append(cur_task)
            cur_sec = new_sec("(概述)")
            cur_task["sections"].append(cur_sec)
            cur_sub = None
        return cur_sec

    for kind, obj in nodes:
        if kind == 't':
            rows = table_to_rows(obj)
            sec = ensure_sec()
            if cur_sub is not None:
                cur_sub["tables"].append(rows)
            else:
                sec["tables"].append(rows)
            continue
        p = obj
        text = clean(p.text)
        if not text: continue
        style = p.style.name if p.style else ''
        hl = hlevel(style)

        m_task = RE_TASK.match(text)
        if m_task:
            tid = f"{m_task.group(1)}.{m_task.group(2)}"
            cur_task = new_task(tid, text)
            tasks.append(cur_task); cur_sec=None; cur_sub=None
            continue
        if hl == 1 and text.startswith("项目"):
            continue
        if hl and ('导学' in text or '智能体互动' in text or is_skip_title(text) or RE_AI.match(text)):
            if cur_task is not None:
                cur_task.setdefault("preamble", []).append(text)
            continue

        # 新单元：编号 一、二、  OR 描述性 Heading 2
        is_numbered_sec = RE_SEC.match(text) is not None
        is_head_sec = (hl == 2) and not is_skip_title(text) and (RE_SUB.match(text) is None) and (RE_NUM.match(text) is None)
        if is_numbered_sec or is_head_sec:
            if cur_task is None:
                cur_task = new_task("0.0","导论"); tasks.append(cur_task)
            cur_sec = new_sec(text); cur_task["sections"].append(cur_sec); cur_sub=None
            continue

        # 子节：Heading3 / Heading4 / （一） / 1. 2.
        is_sub = (hl in (3,4)) or (RE_SUB.match(text) is not None) or (RE_NUM.match(text) is not None)
        if is_sub:
            if cur_sec is None:
                # 无父节：直接作为单元（避免占位爆炸）
                if cur_task is None:
                    cur_task = new_task("0.0","导论"); tasks.append(cur_task)
                cur_sec = new_sec(text); cur_task["sections"].append(cur_sec); cur_sub=None
            else:
                if cur_sub is not None and cur_sub["title"]=="(概述)" and not cur_sub["paras"] and not cur_sub["tables"]:
                    cur_sub["title"] = text
                else:
                    cur_sub = new_sub(text); cur_sec["subs"].append(cur_sub)
            continue

        # 普通段落
        sec = ensure_sec()
        if cur_sub is not None:
            cur_sub["paras"].append(text)
        else:
            sec["paras"].append(text)
    return tasks

def main():
    all_out = {}
    for k,(fn,cid,title) in FILES.items():
        p = os.path.join(SRC, fn)
        tasks = extract(p)
        tasks = [t for t in tasks if t["id"] != "0.0"]  # 丢弃导论
        print(f"\n=== {k} {title} (courseId={cid}) ===")
        print(f"  任务(chapter)={len(tasks)}")
        for t in tasks:
            secs = t["sections"]
            nsub = sum(len(s["subs"]) for s in secs)
            print(f"    任务{t['id']}: 节(单元)={len(secs)} 子节={nsub}  | {t['title'][:34]}")
        all_out[cid] = {"title": title, "tasks": tasks}
    with open(r"D:\workbuddy\chain_supply\_textbook\projects_structure.json","w",encoding="utf-8") as f:
        json.dump(all_out, f, ensure_ascii=False, indent=2)
    print("\n结构已写入 projects_structure.json")

if __name__ == "__main__":
    main()
