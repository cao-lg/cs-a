# -*- coding: utf-8 -*-
"""分析 初稿6.18 中 5 个项目 docx 的结构：大纲层级 + 段落/表格数。"""
import os, re, json
from docx import Document

SRC = r"D:\workbuddy\chain_supply\初稿6.18"
FILES = {
    "p1": "项目一：开启数据驱动之旅_ 第二版.docx",
    "p2": "项目二：商品选品数据分析-修改稿.docx",
    "p3": "项目三：营销推广数据分析 样章3.0.docx",
    "p5": "项目五：客户服务数据分析0501.docx",
    "p6": "项目六：物流履约数据分析3.0.docx",
}

# 大纲行识别：标题样式 或 常见编号模式
HEAD_PATTERNS = [
    re.compile(r'^项目\s*[一二三四五六七八九十\d]'),
    re.compile(r'^任务\s*\d'),
    re.compile(r'^\d+\.\d+\s'),          # 1.1
    re.compile(r'^[一二三四五六七八九十]+、'),  # 一、
    re.compile(r'^（?[一二三四五六七八九十]+）?[\.\s]'),  # （一）
    re.compile(r'^\d+[\.\、]\s'),
]
def looks_head(text, style):
    if not text:
        return False
    t = text.strip()
    if style and ('Heading' in style or '标题' in style):
        return True
    for p in HEAD_PATTERNS:
        if p.match(t):
            return True
    return False

def analyze(path, key):
    doc = Document(path)
    paras = doc.paragraphs
    # 大纲
    outline = []
    n_text = 0
    n_tables = len(doc.tables)
    for p in paras:
        txt = p.text.strip()
        if not txt:
            continue
        n_text += 1
        if looks_head(txt, p.style.name if p.style else ''):
            outline.append((p.style.name if p.style else '', txt[:80]))
    return {
        "key": key,
        "file": os.path.basename(path),
        "n_paragraphs": n_text,
        "n_tables": n_tables,
        "n_outline": len(outline),
        "outline": outline[:120],
    }

if __name__ == "__main__":
    out = {}
    for k, fn in FILES.items():
        p = os.path.join(SRC, fn)
        if not os.path.exists(p):
            print("MISSING", p); continue
        info = analyze(p, k)
        out[k] = info
        print("="*70)
        print(f"{k}  {fn}")
        print(f"  段落={info['n_paragraphs']}  表格={info['n_tables']}  大纲行={info['n_outline']}")
        print("  --- 大纲前若干 ---")
        for st, t in info["outline"][:60]:
            print(f"    [{st}] {t}")
    with open(r"D:\workbuddy\chain_supply\_textbook\docx_analysis.json","w",encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print("\n分析已写入 docx_analysis.json")
